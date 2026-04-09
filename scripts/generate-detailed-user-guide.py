from __future__ import annotations

import json
import os
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path
from urllib.request import urlopen

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from playwright.sync_api import Page, sync_playwright


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
IMAGE_DIR = DOCS_DIR / "detailed-user-guide-assets"
OUTPUT_DOCX = DOCS_DIR / "2gold-huong-dan-nghiep-vu-chi-tiet.docx"
DEMO_META_PATH = ROOT / "tmp" / "doc-data" / "demo-meta.json"
CHROME_PATH = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
DEMO_PORT = "3100"
BASE_URL = f"http://127.0.0.1:{DEMO_PORT}"


def run_command(command: list[str]) -> None:
    subprocess.run(command, cwd=ROOT, check=True)


def wait_for_http(url: str, timeout_seconds: int = 30) -> None:
    deadline = time.time() + timeout_seconds
    last_error = None
    while time.time() < deadline:
        try:
            with urlopen(url, timeout=5) as response:
                if response.status == 200:
                    return
        except Exception as error:  # noqa: BLE001
            last_error = error
            time.sleep(0.5)
    raise RuntimeError(f"Khong the truy cap {url}: {last_error}")


def start_demo_server() -> subprocess.Popen[str]:
    env = os.environ.copy()
    env["DOC_DEMO_PORT"] = DEMO_PORT
    process = subprocess.Popen(
        ["node", "scripts/doc-demo-server.cjs"],
        cwd=ROOT,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    ready = False
    output_lines: list[str] = []
    deadline = time.time() + 30
    while time.time() < deadline:
        line = process.stdout.readline() if process.stdout else ""
        if line:
            output_lines.append(line.rstrip())
            if "DOC_DEMO_SERVER_READY" in line:
                ready = True
                break
        if process.poll() is not None:
            break

    if not ready:
        try:
            process.terminate()
        except Exception:  # noqa: BLE001
            pass
        raise RuntimeError("Khong the khoi dong doc demo server:\n" + "\n".join(output_lines))

    wait_for_http(f"{BASE_URL}/login")
    return process


def load_demo_meta() -> dict:
    deadline = time.time() + 10
    while time.time() < deadline:
        if DEMO_META_PATH.exists():
            return json.loads(DEMO_META_PATH.read_text(encoding="utf-8"))
        time.sleep(0.2)
    raise RuntimeError("Khong tim thay file demo-meta.json")


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_image(document: Document, title: str, image_path: Path) -> None:
    document.add_paragraph(title)
    document.add_picture(str(image_path), width=Inches(6.5))


def highlight(page: Page, selector: str) -> None:
    page.eval_on_selector(
        selector,
        """(element) => {
          element.style.outline = '4px solid #ff6b6b';
          element.style.outlineOffset = '4px';
          element.style.borderRadius = '8px';
          element.scrollIntoView({ behavior: 'instant', block: 'center' });
        }""",
    )


def clear_highlights(page: Page) -> None:
    page.evaluate(
        """() => {
          document.querySelectorAll('*').forEach((element) => {
            if (element instanceof HTMLElement && element.style.outline === '4px solid rgb(255, 107, 107)') {
              element.style.outline = '';
              element.style.outlineOffset = '';
              element.style.borderRadius = '';
            }
          });
        }"""
    )


def capture_screenshots(meta: dict) -> dict[str, Path]:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    shots: dict[str, Path] = {}

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(
            executable_path=str(CHROME_PATH),
            headless=True,
            args=["--disable-gpu", "--window-size=1440,1400"],
        )
        context = browser.new_context(viewport={"width": 1440, "height": 1400}, locale="vi-VN")
        page = context.new_page()

        def save(name: str) -> None:
            output_path = IMAGE_DIR / f"{name}.png"
            page.screenshot(path=str(output_path), full_page=True)
            shots[name] = output_path

        page.goto(f"{BASE_URL}/login", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("login")

        page.fill('input[name="Username"]', "admin")
        page.fill('input[name="Password"]', "secret")
        page.click('button[type="submit"]')
        page.wait_for_url("**/Installment/Index/", timeout=15000)
        page.wait_for_timeout(1200)

        page.goto(f"{BASE_URL}/Installment/Index/", wait_until="domcontentloaded")
        page.wait_for_selector("#installmentTableBody tr[data-row-id]", timeout=10000)
        save("installment-list")
        highlight(page, "#btnImportExcel")
        page.wait_for_timeout(500)
        save("installment-import")
        clear_highlights(page)
        page.click(".js-delete-installment")
        page.wait_for_selector(".swal2-popup", timeout=5000)
        save("installment-delete")
        page.locator(".swal2-cancel").click()
        page.wait_for_timeout(300)
        page.goto(f"{BASE_URL}/Installment/Create", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("installment-create")
        page.goto(f"{BASE_URL}/Installment/Edit/{meta['installments'][0]}", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("installment-edit")

        page.goto(f"{BASE_URL}/Shop/Index/", wait_until="domcontentloaded")
        page.wait_for_selector("#shopTableBody tr[data-row-id]", timeout=10000)
        save("shop-list")
        page.click(".js-delete-shop")
        page.wait_for_selector(".swal2-popup", timeout=5000)
        save("shop-delete")
        page.locator(".swal2-cancel").click()
        page.wait_for_timeout(300)
        page.goto(f"{BASE_URL}/Shop/Create", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("shop-create")
        page.goto(f"{BASE_URL}/Shop/Create?ShopID={meta['shops'][0]}", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("shop-edit")

        page.goto(f"{BASE_URL}/Staff/Index/", wait_until="domcontentloaded")
        page.wait_for_selector("#staffTableBody tr[data-row-id]", timeout=10000)
        save("staff-list")
        page.click(".js-delete-staff")
        page.wait_for_selector(".swal2-popup", timeout=5000)
        save("staff-delete")
        page.locator(".swal2-cancel").click()
        page.wait_for_timeout(300)
        page.goto(f"{BASE_URL}/Staff/Create", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("staff-create")
        page.goto(f"{BASE_URL}/Staff/Create?StaffID={meta['staff']['staffAId']}", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("staff-edit")
        page.goto(f"{BASE_URL}/Staff/PermissionStaff/?StaffID={meta['staff']['staffAId']}", wait_until="domcontentloaded")
        page.wait_for_selector("form", timeout=10000)
        save("staff-permission")

        page.goto(f"{BASE_URL}/History/", wait_until="domcontentloaded")
        page.wait_for_selector(".history-table", timeout=10000)
        save("history-list")

        browser.close()

    return shots


def build_document(shots: dict[str, Path], meta: dict) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = document.add_heading("TÃ€I LIá»†U HÆ¯á»šNG DáºªN NGHIá»†P Vá»¤ CHI TIáº¾T Há»† THá»NG 2GOLD", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"NgÃ y táº¡o: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

    document.add_paragraph(
        "TÃ i liá»‡u nÃ y mÃ´ táº£ cÃ¡ch sá»­ dá»¥ng há»‡ thá»‘ng theo gÃ³c nhÃ¬n nghiá»‡p vá»¥. Má»—i tab bao gá»“m má»¥c Ä‘Ã­ch sá»­ dá»¥ng, danh sÃ¡ch nÃºt chá»©c nÄƒng, luá»“ng thao tÃ¡c, lÆ°u Ã½ phÃ¢n quyá»n vÃ  áº£nh minh há»a cho cÃ¡c thao tÃ¡c chÃ­nh."
    )

    document.add_heading("1. Quy Æ°á»›c chung", level=1)
    add_bullet(document, "Admin: cÃ³ quyá»n quáº£n trá»‹ toÃ n há»‡ thá»‘ng, quáº£n lÃ½ nhiá»u cá»­a hÃ ng, phÃ¢n quyá»n nhÃ¢n viÃªn, thao tÃ¡c toÃ n bá»™ dá»¯ liá»‡u.")
    add_bullet(document, "NhÃ¢n viÃªn: chá»‰ thao tÃ¡c trong pháº¡m vi cá»­a hÃ ng vÃ  module Ä‘Æ°á»£c cáº¥p quyá»n.")
    add_bullet(document, "CÃ¡c báº£ng dá»¯ liá»‡u Ä‘á»u há»— trá»£ chá»n nhiá»u dÃ²ng Ä‘á»ƒ thao tÃ¡c hÃ ng loáº¡t khi tÃ i khoáº£n cÃ³ quyá»n phÃ¹ há»£p.")
    add_bullet(document, f"MÃ´i trÆ°á»ng minh há»a Ä‘ang dÃ¹ng 2 cá»­a hÃ ng máº«u vÃ  {len(meta['installments'])} há»£p Ä‘á»“ng tráº£ gÃ³p máº«u.")

    document.add_heading("2. Trang Ä‘Äƒng nháº­p", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: xÃ¡c thá»±c tÃ i khoáº£n trÆ°á»›c khi vÃ o há»‡ thá»‘ng.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "Ã” TÃªn Ä‘Äƒng nháº­p.")
    add_bullet(document, "Ã” Máº­t kháº©u.")
    add_bullet(document, "Checkbox Ghi nhá»› Ä‘Äƒng nháº­p.")
    add_bullet(document, "NÃºt ÄÄƒng nháº­p.")
    document.add_paragraph("Luá»“ng thao tÃ¡c:")
    add_number(document, "Má»Ÿ trang Ä‘Äƒng nháº­p táº¡i /login.")
    add_number(document, "Nháº­p tÃ i khoáº£n vÃ  máº­t kháº©u.")
    add_number(document, "Nháº¥n ÄÄƒng nháº­p Ä‘á»ƒ vÃ o trang tá»•ng quan.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Quyá»n thá»±c táº¿ cá»§a ngÆ°á»i dÃ¹ng sáº½ Ä‘Æ°á»£c Ã¡p sau khi Ä‘Äƒng nháº­p thÃ nh cÃ´ng.")
    add_bullet(document, "Náº¿u admin thay Ä‘á»•i phÃ¢n quyá»n cho nhÃ¢n viÃªn, ngÆ°á»i Ä‘Ã³ cáº§n Ä‘Äƒng nháº­p láº¡i Ä‘á»ƒ nháº­n quyá»n má»›i.")
    add_image(document, "áº¢nh máº«u trang Ä‘Äƒng nháº­p:", shots["login"])

    document.add_heading("3. Tab Tráº£ gÃ³p", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: quáº£n lÃ½ danh sÃ¡ch há»£p Ä‘á»“ng tráº£ gÃ³p, nháº­p dá»¯ liá»‡u Excel, cáº­p nháº­t tráº¡ng thÃ¡i vÃ  theo dÃµi doanh thu theo cá»­a hÃ ng.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "ThÃªm má»›i: má»Ÿ form táº¡o há»£p Ä‘á»“ng má»›i.")
    add_bullet(document, "Láº¥y dá»¯ liá»‡u: táº£i láº¡i danh sÃ¡ch theo bá»™ lá»c hiá»‡n táº¡i.")
    add_bullet(document, "Nháº­p Excel: náº¡p file Excel vÃ o danh sÃ¡ch vÃ  lÆ°u database.")
    add_bullet(document, "Sá»­a: cáº­p nháº­t tá»«ng há»£p Ä‘á»“ng.")
    add_bullet(document, "XÃ³a: xÃ³a tá»«ng há»£p Ä‘á»“ng.")
    add_bullet(document, "Äá»•i tráº¡ng thÃ¡i: cáº­p nháº­t tráº¡ng thÃ¡i cho nhiá»u há»£p Ä‘á»“ng Ä‘Ã£ chá»n.")
    add_bullet(document, "Export CSV: xuáº¥t cÃ¡c dÃ²ng Ä‘ang chá»n.")
    document.add_paragraph("Luá»“ng thao tÃ¡c thÃªm má»›i:")
    add_number(document, "VÃ o tab Tráº£ gÃ³p > Danh sÃ¡ch.")
    add_number(document, "Nháº¥n ThÃªm má»›i Ä‘á»ƒ má»Ÿ form táº¡o há»£p Ä‘á»“ng.")
    add_number(document, "Chá»n cá»­a hÃ ng, nháº­p thÃ´ng tin khÃ¡ch, IMEI, gÃ³i vay, doanh thu vÃ  tráº¡ng thÃ¡i.")
    add_number(document, "Nháº¥n LÆ°u má»›i Ä‘á»ƒ ghi dá»¯ liá»‡u vÃ o database.")
    document.add_paragraph("Luá»“ng thao tÃ¡c sá»­a:")
    add_number(document, "Trong danh sÃ¡ch, tÃ¬m Ä‘Ãºng há»£p Ä‘á»“ng cáº§n sá»­a.")
    add_number(document, "Nháº¥n Sá»­a á»Ÿ cuá»‘i dÃ²ng.")
    add_number(document, "Cáº­p nháº­t cÃ¡c trÆ°á»ng cáº§n thay Ä‘á»•i vÃ  lÆ°u láº¡i.")
    document.add_paragraph("Luá»“ng thao tÃ¡c xÃ³a:")
    add_number(document, "Nháº¥n XÃ³a táº¡i dÃ²ng cáº§n xÃ³a.")
    add_number(document, "XÃ¡c nháº­n trong há»™p thoáº¡i cáº£nh bÃ¡o.")
    add_number(document, "Há»‡ thá»‘ng xÃ³a báº£n ghi vÃ  ghi lá»‹ch sá»­ thao tÃ¡c.")
    document.add_paragraph("Luá»“ng thao tÃ¡c import Excel:")
    add_number(document, "Nháº¥n Nháº­p Excel táº¡i mÃ n hÃ¬nh danh sÃ¡ch.")
    add_number(document, "Chá»n file Excel Ä‘Ãºng cáº¥u trÃºc cá»™t cá»§a há»‡ thá»‘ng.")
    add_number(document, "Há»‡ thá»‘ng Ä‘á»c file, chuáº©n hÃ³a ngÃ y vÃ  sá»‘ tiá»n rá»“i lÆ°u vÃ o database.")
    add_number(document, "Táº£i láº¡i danh sÃ¡ch Ä‘á»ƒ kiá»ƒm tra dá»¯ liá»‡u sau import.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Admin cÃ³ thá»ƒ thao tÃ¡c dá»¯ liá»‡u nhiá»u cá»­a hÃ ng.")
    add_bullet(document, "NhÃ¢n viÃªn chá»‰ xem vÃ  sá»­a dá»¯ liá»‡u thuá»™c cá»­a hÃ ng Ä‘Æ°á»£c phÃ¢n quyá»n.")
    add_bullet(document, "Import Excel pháº£i gáº¯n vá»›i má»™t cá»­a hÃ ng cá»¥ thá»ƒ.")
    add_image(document, "áº¢nh danh sÃ¡ch Tráº£ gÃ³p:", shots["installment-list"])
    add_image(document, "áº¢nh thao tÃ¡c ThÃªm má»›i Tráº£ gÃ³p:", shots["installment-create"])
    add_image(document, "áº¢nh thao tÃ¡c Sá»­a Tráº£ gÃ³p:", shots["installment-edit"])
    add_image(document, "áº¢nh thao tÃ¡c XÃ³a Tráº£ gÃ³p:", shots["installment-delete"])
    add_image(document, "áº¢nh thao tÃ¡c Import Excel Tráº£ gÃ³p:", shots["installment-import"])

    document.add_heading("4. Tab Cá»­a hÃ ng", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: quáº£n lÃ½ danh sÃ¡ch cá»­a hÃ ng, thÃ´ng tin liÃªn há»‡, vá»‘n Ä‘áº§u tÆ° vÃ  tráº¡ng thÃ¡i hoáº¡t Ä‘á»™ng.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "ThÃªm má»›i: táº¡o cá»­a hÃ ng má»›i.")
    add_bullet(document, "Láº¥y dá»¯ liá»‡u: táº£i láº¡i danh sÃ¡ch cá»­a hÃ ng.")
    add_bullet(document, "Sá»­a: cáº­p nháº­t thÃ´ng tin cá»­a hÃ ng.")
    add_bullet(document, "XÃ³a: xÃ³a cá»­a hÃ ng.")
    add_bullet(document, "Äá»•i tráº¡ng thÃ¡i: Ä‘á»•i tráº¡ng thÃ¡i hoáº¡t Ä‘á»™ng hÃ ng loáº¡t.")
    add_bullet(document, "Export CSV: xuáº¥t cÃ¡c cá»­a hÃ ng Ä‘ang chá»n.")
    document.add_paragraph("Luá»“ng thao tÃ¡c thÃªm má»›i:")
    add_number(document, "VÃ o tab Cá»­a hÃ ng > Danh sÃ¡ch.")
    add_number(document, "Nháº¥n ThÃªm má»›i Ä‘á»ƒ má»Ÿ form.")
    add_number(document, "Nháº­p tÃªn cá»­a hÃ ng, Ä‘iá»‡n thoáº¡i, Ä‘áº¡i diá»‡n, vá»‘n Ä‘áº§u tÆ° vÃ  ngÃ y táº¡o.")
    add_number(document, "Chá»n Ä‘á»‹a chá»‰ theo tá»‰nh/thÃ nh, quáº­n/huyá»‡n tÆ°Æ¡ng thÃ­ch vÃ  phÆ°á»ng/xÃ£.")
    add_number(document, "Nháº¥n LÆ°u má»›i Ä‘á»ƒ ghi dá»¯ liá»‡u.")
    document.add_paragraph("Luá»“ng thao tÃ¡c sá»­a:")
    add_number(document, "Nháº¥n Sá»­a táº¡i dÃ²ng cá»­a hÃ ng cáº§n cáº­p nháº­t.")
    add_number(document, "Äiá»u chá»‰nh thÃ´ng tin vÃ  nháº¥n Cáº­p nháº­t.")
    document.add_paragraph("Luá»“ng thao tÃ¡c xÃ³a:")
    add_number(document, "Nháº¥n XÃ³a táº¡i dÃ²ng cá»­a hÃ ng.")
    add_number(document, "XÃ¡c nháº­n trong popup trÆ°á»›c khi há»‡ thá»‘ng xÃ³a dá»¯ liá»‡u.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Admin cÃ³ quyá»n CRUD toÃ n bá»™ cá»­a hÃ ng.")
    add_bullet(document, "NhÃ¢n viÃªn chá»‰ xem trong pháº¡m vi cá»­a hÃ ng Ä‘Æ°á»£c gÃ¡n, khÃ´ng cÃ³ quyá»n xÃ³a náº¿u khÃ´ng Ä‘Æ°á»£c cáº¥p.")
    add_image(document, "áº¢nh danh sÃ¡ch Cá»­a hÃ ng:", shots["shop-list"])
    add_image(document, "áº¢nh thao tÃ¡c ThÃªm má»›i Cá»­a hÃ ng:", shots["shop-create"])
    add_image(document, "áº¢nh thao tÃ¡c Sá»­a Cá»­a hÃ ng:", shots["shop-edit"])
    add_image(document, "áº¢nh thao tÃ¡c XÃ³a Cá»­a hÃ ng:", shots["shop-delete"])

    document.add_heading("5. Tab NhÃ¢n viÃªn", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: quáº£n lÃ½ tÃ i khoáº£n ná»™i bá»™, tráº¡ng thÃ¡i lÃ m viá»‡c vÃ  gáº¯n nhÃ¢n viÃªn vÃ o cá»­a hÃ ng tÆ°Æ¡ng á»©ng.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "ThÃªm má»›i: táº¡o nhÃ¢n viÃªn má»›i.")
    add_bullet(document, "Láº¥y dá»¯ liá»‡u: táº£i láº¡i danh sÃ¡ch nhÃ¢n viÃªn.")
    add_bullet(document, "Sá»­a: cáº­p nháº­t há»“ sÆ¡ nhÃ¢n viÃªn.")
    add_bullet(document, "XÃ³a: xÃ³a nhÃ¢n viÃªn.")
    add_bullet(document, "Äá»•i tráº¡ng thÃ¡i: khÃ³a/má»Ÿ khÃ³a nhiá»u nhÃ¢n viÃªn.")
    add_bullet(document, "Export CSV: xuáº¥t danh sÃ¡ch nhÃ¢n viÃªn Ä‘Ã£ chá»n.")
    add_bullet(document, "Quyá»n: má»Ÿ tab phÃ¢n quyá»n nhÃ¢n viÃªn.")
    document.add_paragraph("Luá»“ng thao tÃ¡c thÃªm má»›i:")
    add_number(document, "VÃ o tab NhÃ¢n viÃªn > Danh sÃ¡ch.")
    add_number(document, "Nháº¥n ThÃªm má»›i.")
    add_number(document, "Nháº­p tÃ i khoáº£n, há» tÃªn, email, Ä‘iá»‡n thoáº¡i, ngÃ y táº¡o vÃ  cá»­a hÃ ng.")
    add_number(document, "Nháº¥n LÆ°u má»›i Ä‘á»ƒ táº¡o tÃ i khoáº£n nhÃ¢n viÃªn ná»™i bá»™.")
    document.add_paragraph("Luá»“ng thao tÃ¡c sá»­a:")
    add_number(document, "Nháº¥n Sá»­a á»Ÿ dÃ²ng nhÃ¢n viÃªn cáº§n cáº­p nháº­t.")
    add_number(document, "Thay Ä‘á»•i thÃ´ng tin vÃ  lÆ°u láº¡i.")
    document.add_paragraph("Luá»“ng thao tÃ¡c xÃ³a:")
    add_number(document, "Nháº¥n XÃ³a á»Ÿ dÃ²ng nhÃ¢n viÃªn.")
    add_number(document, "XÃ¡c nháº­n thao tÃ¡c trong popup.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Admin cÃ³ thá»ƒ xem toÃ n bá»™ danh sÃ¡ch vÃ  phÃ¢n quyá»n.")
    add_bullet(document, "NhÃ¢n viÃªn thÆ°á»ng chá»‰ thao tÃ¡c trong pháº¡m vi cá»­a hÃ ng Ä‘Æ°á»£c cáº¥p.")
    add_bullet(document, "Quyá»n module vÃ  pháº¡m vi cá»­a hÃ ng Ä‘Æ°á»£c chá»‰nh á»Ÿ mÃ n hÃ¬nh PhÃ¢n quyá»n.")
    add_image(document, "áº¢nh danh sÃ¡ch NhÃ¢n viÃªn:", shots["staff-list"])
    add_image(document, "áº¢nh thao tÃ¡c ThÃªm má»›i NhÃ¢n viÃªn:", shots["staff-create"])
    add_image(document, "áº¢nh thao tÃ¡c Sá»­a NhÃ¢n viÃªn:", shots["staff-edit"])
    add_image(document, "áº¢nh thao tÃ¡c XÃ³a NhÃ¢n viÃªn:", shots["staff-delete"])

    document.add_heading("6. Tab PhÃ¢n quyá»n", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: cho phÃ©p Admin cáº¥u hÃ¬nh vai trÃ², module truy cáº­p vÃ  pháº¡m vi cá»­a hÃ ng cá»§a tá»«ng nhÃ¢n viÃªn.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "Chá»n nhÃ¢n viÃªn cáº§n phÃ¢n quyá»n.")
    add_bullet(document, "Chá»n role admin hoáº·c staff.")
    add_bullet(document, "Chá»n cá»­a hÃ ng máº·c Ä‘á»‹nh vÃ  danh sÃ¡ch cá»­a hÃ ng Ä‘Æ°á»£c phÃ©p truy cáº­p.")
    add_bullet(document, "Chá»n module Ä‘Æ°á»£c phÃ©p thao tÃ¡c.")
    add_bullet(document, "LÆ°u cáº¥u hÃ¬nh phÃ¢n quyá»n.")
    document.add_paragraph("Luá»“ng thao tÃ¡c:")
    add_number(document, "VÃ o tab PhÃ¢n quyá»n.")
    add_number(document, "Chá»n nhÃ¢n viÃªn cáº§n cáº¥u hÃ¬nh.")
    add_number(document, "Äáº·t vai trÃ², cá»­a hÃ ng máº·c Ä‘á»‹nh vÃ  cÃ¡c module Ä‘Æ°á»£c truy cáº­p.")
    add_number(document, "LÆ°u láº¡i vÃ  yÃªu cáº§u ngÆ°á»i dÃ¹ng Ä‘Äƒng nháº­p láº¡i Ä‘á»ƒ nháº­n quyá»n má»›i.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Chá»‰ Admin má»›i vÃ o Ä‘Æ°á»£c mÃ n hÃ¬nh nÃ y.")
    add_bullet(document, "Náº¿u role lÃ  admin thÃ¬ tÃ i khoáº£n cÃ³ thá»ƒ Ä‘Æ°á»£c phÃ©p thao tÃ¡c pháº¡m vi rá»™ng hÆ¡n.")
    add_image(document, "áº¢nh mÃ n hÃ¬nh PhÃ¢n quyá»n:", shots["staff-permission"])

    document.add_heading("7. Tab History", level=1)
    document.add_paragraph("Má»¥c Ä‘Ã­ch sá»­ dá»¥ng: theo dÃµi nháº­t kÃ½ truy cáº­p vÃ  lá»‹ch sá»­ CRUD cá»§a ngÆ°á»i dÃ¹ng trong há»‡ thá»‘ng.")
    document.add_paragraph("CÃ¡c nÃºt/chá»©c nÄƒng chÃ­nh:")
    add_bullet(document, "Bá»™ lá»c tá»« khÃ³a.")
    add_bullet(document, "Lá»c theo module, loáº¡i thao tÃ¡c, tÃ i khoáº£n, cá»­a hÃ ng vÃ  ngÃ y.")
    add_bullet(document, "PhÃ¢n trang lá»‹ch sá»­ thao tÃ¡c.")
    document.add_paragraph("Luá»“ng thao tÃ¡c:")
    add_number(document, "Má»Ÿ tab History.")
    add_number(document, "Chá»n bá»™ lá»c phÃ¹ há»£p vá»›i nghiá»‡p vá»¥ cáº§n tra cá»©u.")
    add_number(document, "Nháº¥n Lá»c dá»¯ liá»‡u Ä‘á»ƒ xem nháº­t kÃ½ phÃ¹ há»£p.")
    document.add_paragraph("LÆ°u Ã½ phÃ¢n quyá»n:")
    add_bullet(document, "Admin cÃ³ thá»ƒ xem toÃ n bá»™ lá»‹ch sá»­.")
    add_bullet(document, "NhÃ¢n viÃªn chá»‰ xem Ä‘Æ°á»£c pháº¡m vi log thuá»™c quyá»n cá»§a mÃ¬nh.")
    add_image(document, "áº¢nh mÃ n hÃ¬nh History:", shots["history-list"])

    document.save(OUTPUT_DOCX)


def main() -> int:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    run_command(["cmd", "/c", "npm.cmd", "run", "build"])

    process = start_demo_server()
    try:
        meta = load_demo_meta()
        shots = capture_screenshots(meta)
        build_document(shots, meta)
    finally:
        process.terminate()
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            process.kill()

    print(f"Generated: {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())



