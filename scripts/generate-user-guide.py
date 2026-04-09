from __future__ import annotations

import json
import os
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path
from urllib.request import urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from playwright.sync_api import sync_playwright


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
IMAGE_DIR = DOCS_DIR / "user-guide-assets"
OUTPUT_DOCX = DOCS_DIR / "2gold-huong-dan-su-dung.docx"
DEMO_META_PATH = ROOT / "tmp" / "doc-data" / "demo-meta.json"
CHROME_PATH = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
DEMO_PORT = "3100"
BASE_URL = f"http://127.0.0.1:{DEMO_PORT}"


def run_command(command: list[str]) -> None:
    subprocess.run(command, cwd=ROOT, check=True)


def wait_for_http(url: str, timeout_seconds: int = 30) -> None:
    deadline = time.time() + timeout_seconds
    last_error = None
    while time.time() < deadline:
        try:
            with urlopen(url, timeout=5) as response:
                if response.status == 200:
                    return
        except Exception as error:  # noqa: BLE001
            last_error = error
            time.sleep(0.5)

    raise RuntimeError(f"Khong the truy cap {url}: {last_error}")


def start_demo_server() -> subprocess.Popen[str]:
    env = os.environ.copy()
    env["DOC_DEMO_PORT"] = DEMO_PORT
    process = subprocess.Popen(
        ["node", "scripts/doc-demo-server.cjs"],
        cwd=ROOT,
        env=env,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )

    ready = False
    output_lines: list[str] = []
    deadline = time.time() + 30
    while time.time() < deadline:
        line = process.stdout.readline() if process.stdout else ""
        if line:
            output_lines.append(line.rstrip())
            if "DOC_DEMO_SERVER_READY" in line:
                ready = True
                break
        if process.poll() is not None:
            break

    if not ready:
        try:
            process.terminate()
        except Exception:  # noqa: BLE001
            pass
        raise RuntimeError("Khong the khoi dong doc demo server:\n" + "\n".join(output_lines))

    wait_for_http(f"{BASE_URL}/login")
    return process


def load_demo_meta() -> dict:
    deadline = time.time() + 10
    while time.time() < deadline:
        if DEMO_META_PATH.exists():
            return json.loads(DEMO_META_PATH.read_text(encoding="utf-8"))
        time.sleep(0.2)
    raise RuntimeError("Khong tim thay file demo-meta.json")


def capture_screenshots(meta: dict) -> dict[str, Path]:
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)
    shots: dict[str, Path] = {}

    with sync_playwright() as playwright:
        browser = playwright.chromium.launch(
            executable_path=str(CHROME_PATH),
            headless=True,
            args=["--disable-gpu", "--window-size=1440,1200"],
        )
        context = browser.new_context(viewport={"width": 1440, "height": 1200}, locale="vi-VN")
        page = context.new_page()

        def capture(name: str, url: str, selector: str | None = None, timeout_ms: int = 1500) -> None:
            page.goto(url, wait_until="domcontentloaded")
            if selector:
                page.wait_for_selector(selector, timeout=10000)
            page.wait_for_timeout(timeout_ms)
            output_path = IMAGE_DIR / f"{name}.png"
            page.screenshot(path=str(output_path), full_page=True)
            shots[name] = output_path

        capture("login", f"{BASE_URL}/login", "form")

        page.fill('input[name="Username"]', "admin")
        page.fill('input[name="Password"]', "secret")
        page.click('button[type="submit"]')
        page.wait_for_url("**/Installment/Index/", timeout=15000)
        page.wait_for_timeout(1200)
        shots["home"] = IMAGE_DIR / "home.png"
        page.screenshot(path=str(shots["home"]), full_page=True)

        capture("installment-index", f"{BASE_URL}/Installment/Index/", "#installmentTableBody tr[data-row-id]")
        capture("installment-create", f"{BASE_URL}/Installment/Create", "form")
        capture("shop-index", f"{BASE_URL}/Shop/Index/", "#shopTableBody tr[data-row-id]")
        capture("shop-create", f"{BASE_URL}/Shop/Create", "form")
        capture("staff-index", f"{BASE_URL}/Staff/Index/", "#staffTableBody tr[data-row-id]")
        capture(
            "staff-permission",
            f"{BASE_URL}/Staff/PermissionStaff/?StaffID={meta['staff']['staffAId']}",
            "form",
        )
        capture("history", f"{BASE_URL}/History/", ".history-table")

        browser.close()

    return shots


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def build_document(shots: dict[str, Path], meta: dict) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    title = document.add_heading("TÃ€I LIá»†U HÆ¯á»šNG DáºªN Sá»¬ Dá»¤NG Há»† THá»NG 2GOLD Ná»˜I Bá»˜", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"NgÃ y táº¡o tÃ i liá»‡u: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

    document.add_paragraph(
        "TÃ i liá»‡u nÃ y Ä‘Æ°á»£c táº¡o tá»± Ä‘á»™ng tá»« giao diá»‡n hiá»‡n táº¡i cá»§a há»‡ thá»‘ng, sá»­ dá»¥ng mÃ´i trÆ°á»ng demo ná»™i bá»™ cÃ³ dá»¯ liá»‡u máº«u Ä‘á»ƒ minh há»a cÃ¡c tab vÃ  chá»©c nÄƒng chÃ­nh."
    )

    document.add_heading("1. ThÃ´ng tin demo", level=1)
    add_bullet(document, f"URL demo dÃ¹ng Ä‘á»ƒ chá»¥p áº£nh: {BASE_URL}")
    add_bullet(document, "TÃ i khoáº£n quáº£n trá»‹ demo: admin / secret")
    add_bullet(document, "TÃ i khoáº£n nhÃ¢n viÃªn demo: staff01 / secret")
    add_bullet(document, f"Sá»‘ cá»­a hÃ ng máº«u: {len(meta['shops'])}")
    add_bullet(document, f"Sá»‘ há»£p Ä‘á»“ng tráº£ gÃ³p máº«u: {len(meta['installments'])}")

    sections = [
        {
            "title": "2. Trang dang nh?p",
            "url": "/login",
            "image": shots["login"],
            "features": [
                "Cho phép ngu?i dùng dang nh?p vào h? th?ng b?ng tài kho?n n?i b? ho?c tài kho?n du?c d?ng b? quy?n.",
                "G?m các tru?ng tên dang nh?p, m?t kh?u và tùy ch?n ghi nh? dang nh?p.",
                "Là di?m xu?t phát tru?c khi truy c?p các module nhu Tr? góp, C?a hàng, Nhân viên, Phân quy?n và L?ch s? thao tác.",
            ],
        },
        {
            "title": "3. Tab Tr? góp - Danh sách",
            "url": "/Installment/Index/",
            "image": shots["installment-index"],
            "features": [
                "Hi?n th? dashboard th?ng kê, b?ng danh sách h?p d?ng kèm b? l?c và các thao tác CRUD.",
                "H? tr? nh?p Excel, ch?n nhi?u dòng, xoá nhi?u b?n ghi, d?i tr?ng thái và xu?t CSV cho t?p d? li?u.",
                "D? li?u ràng bu?c theo c?a hàng và phân quy?n c?a ngu?i dùng dang dang nh?p.",
            ],
        },
        {
            "title": "4. Tab Tr? góp - Thêm m?i",
            "url": "/Installment/Create",
            "image": shots["installment-create"],
            "features": [
                "Dùng d? t?o m?i m?t h?p d?ng tr? góp n?i b?.",
                "Các tru?ng ngày s? d?ng b? ch?n date; các tru?ng s? ti?n du?c chu?n hóa tru?c khi luu.",
                "Ngu?i qu?n lý có th? gán h?p d?ng vào c?a hàng c? th? tru?c khi luu d? li?u.",
            ],
        },
        {
            "title": "5. Tab C?a hàng - Danh sách",
            "url": "/Shop/Index/",
            "image": shots["shop-index"],
            "features": [
                "Hi?n th? danh sách c?a hàng, dashboard v?n d?u tu và tr?ng thái ho?t d?ng.",
                "H? tr? l?c theo t? khóa, tr?ng thái, phân trang và thao tác nhi?u dòng nhu xóa ho?c d?i tr?ng thái.",
                "Quy?n qu?n tr? quy?t d?nh kh? nang t?o m?i, s?a ho?c xóa c?a hàng.",
            ],
        },
        {
            "title": "6. Tab C?a hàng - Thêm m?i",
            "url": "/Shop/Create",
            "image": shots["shop-create"],
            "features": [
                "Dùng d? t?o m?i ho?c c?p nh?t thông tin c?a hàng.",
                "Tru?ng d?a ch? h? tr? ch?n t?nh/thành, qu?n/huy?n và phu?ng/xã t? d? li?u d?a lý.",
                "Tru?ng v?n d?u tu nh?p theo d?nh d?ng ti?n Vi?t Nam và du?c chu?n hóa khi luu.",
            ],
        },
        {
            "title": "7. Tab Nhân viên - Danh sách",
            "url": "/Staff/Index/",
            "image": shots["staff-index"],
            "features": [
                "Hi?n th? danh sách nhân viên theo c?a hàng, tr?ng thái và vai trò.",
                "H? tr? CRUD, ch?n nhi?u dòng, xóa nhi?u, d?i tr?ng thái và xu?t CSV.",
                "Nhân viên thu?ng ch? th?y d? li?u c?a hàng du?c phân quy?n; qu?n tr? có th? xem toàn b?.",
            ],
        },
        {
            "title": "8. Tab Phân quy?n nhân viên",
            "url": f"/Staff/PermissionStaff/?StaffID={meta['staff']['staffAId']}",
            "image": shots["staff-permission"],
            "features": [
                "Dành cho Admin c?u hình vai trò, ph?m vi c?a hàng và quy?n module cho t?ng nhân viên.",
                "Cho phép xác d?nh nhân viên du?c phép thao tác module nào và truy c?p c?a hàng nào.",
                "Sau khi d?i quy?n, ngu?i dùng c?n dang xu?t và dang nh?p l?i d? nh?n quy?n m?i.",
            ],
        },
        {
            "title": "9. Tab L?ch s? thao tác",
            "url": "/History/",
            "image": shots["history"],
            "features": [
                "Ghi nh?n l?ch s? dang nh?p, truy c?p trang, CRUD, import và phân quy?n c?a toàn b? ngu?i dùng.",
                "Có b? l?c theo module, thao tác, tài kho?n, c?a hàng, kho?ng ngày và s? dòng m?i trang.",
                "Phù h?p d? ki?m tra thao tác phát sinh trong h? th?ng và ph?c v? truy v?t n?i b?.",
            ],
        },
    ]

    for section in sections:
        document.add_section(WD_SECTION.CONTINUOUS)
        document.add_heading(section["title"], level=1)
        url_paragraph = document.add_paragraph()
        url_paragraph.add_run("ÄÆ°á»ng dáº«n: ").bold = True
        url_paragraph.add_run(section["url"])
        for feature in section["features"]:
            add_bullet(document, feature)
        document.add_paragraph("áº¢nh minh há»a giao diá»‡n hiá»‡n táº¡i:")
        document.add_picture(str(section["image"]), width=Inches(6.5))

    document.save(OUTPUT_DOCX)


def main() -> int:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    print("Building project...")
    run_command(["cmd", "/c", "npm.cmd", "run", "build"])

    print("Starting demo server...")
    process = start_demo_server()
    try:
        meta = load_demo_meta()
        print("Capturing screenshots...")
        shots = capture_screenshots(meta)
        print("Generating Word document...")
        build_document(shots, meta)
    finally:
        process.terminate()
        try:
            process.wait(timeout=10)
        except subprocess.TimeoutExpired:
            process.kill()

    print(f"Generated: {OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())



